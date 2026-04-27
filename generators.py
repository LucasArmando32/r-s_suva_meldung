"""
Generatoren fuer die SUVA-Anmeldungs-Dokumente.
Erstellt: Sanierungsplan.docx, Luftbilanz.xlsx, Alarmliste.docx
"""
import os
import shutil
from datetime import datetime

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl

# =========================================================================
# Hilfsfunktionen
# =========================================================================

CHECKED = "\u2612"
UNCHECKED = "\u2610"
RS_RED = RGBColor(0xC8, 0x10, 0x2E)


def format_date(iso_date: str) -> str:
    """ISO-Datum (YYYY-MM-DD) -> DD.MM.YYYY"""
    if not iso_date:
        return ""
    try:
        return datetime.strptime(iso_date, "%Y-%m-%d").strftime("%d.%m.%Y")
    except ValueError:
        return iso_date


def replace_in_paragraph(paragraph, replacements: dict):
    """Ersetzt Text in einem Paragraphen. Kombiniert Runs."""
    if not paragraph.runs:
        return
    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text
    changed = False
    for old, new in replacements.items():
        if old in new_text:
            new_text = new_text.replace(old, new)
            changed = True
    if changed:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def replace_in_document(doc: Document, replacements: dict):
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)


def set_cell_bg(cell, hex_color: str):
    """Setzt die Hintergrundfarbe einer Tabellenzelle."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# =========================================================================
# Sanierungsplan Generator
# =========================================================================

def generate_sanierungsplan(data: dict, template_path: str, output_path: str):
    """Generiert den Sanierungsplan aus dem gewaehlten Template."""
    doc = Document(template_path)

    personal_lines = [l.strip() for l in data.get("personal", "").split("\n") if l.strip()]
    while len(personal_lines) < 5:
        personal_lines.append("")
    p_line1 = f"{personal_lines[0]}, \t\t{personal_lines[1]}"
    p_line2 = f"{personal_lines[2]}, \t{personal_lines[3]}"
    p_line3 = personal_lines[4]

    datum_bearb = format_date(data.get("datum_bearbeitung"))
    san_von = format_date(data.get("sanierung_von"))
    san_bis = format_date(data.get("sanierung_bis"))
    besprech = format_date(data.get("besprechung_datum"))
    sanierung_zeitraum = f"{san_von} bis {san_bis}" if san_bis else san_von

    # Gemeinsame Replacements fuer BEIDE Templates (EKAS/Factsheet)
    # Der Text "Tulpenweg" ist im Factsheet-Template vorhanden,
    # "Steinrieselstrasse" im EKAS-Template. Beide werden ersetzt.
    replacements = {
        # === FACTSHEET Template (Tulpenweg-Daten) ===
        "MFH Badewanne Sanierung Tulpenweg": data.get("baustelle_objekt", ""),
        "TulpenWeg 105": data.get("baustelle_adresse", ""),
        "3098 Köniz": data.get("baustelle_plz_ort", ""),
        "Pfister Robert AG": data.get("auftraggeber_name", ""),
        "Lilienweg 46": data.get("auftraggeber_adresse", ""),
        "Helmut Schafroth": data.get("bauherr_firma", ""),
        "Tulpenweg 105": data.get("bauherr_adresse", ""),
        "sr@ru-energie.ch": data.get("bauherr_email", ""),
        "Hinterkappelen 22.08.2023": f"{data.get('ort_bearbeitung', '')} {datum_bearb}",
        "Hinterkappelen, 29.08.2023": f"{data.get('ort_bearbeitung', '')}, {datum_bearb}",
        "Asbesthaltiger Badewanne (Wandfliesen)": data.get("schadstoff_art", ""),
        ">1 m2": data.get("schadstoff_menge", ""),
        "Badezimmer -> 1 Kammerschleuse": data.get("schadstoff_lage", ""),
        "Cintia Reitmann 756.5257.7207.03, \t\tWerner Neuhaus 756.5596.8850.64": p_line1,
        "Pablo Cabrera Harisgarat 756.3394.1133.27, \tJose Alberto Chica Briones 756.7932.7946.58": p_line2,
        "Luis Enrique Silva Ccanto 756.2154.4658.05": p_line3,
        "Sanierung vom 31.08.2023": f"Sanierung vom {sanierung_zeitraum}",
        "Besprechung/Baustellensitzung/Begehung vom 02.06.2023":
            f"Besprechung/Baustellensitzung/Begehung vom {besprech}" if besprech else
            "Besprechung/Baustellensitzung/Begehung vom",

        # === EKAS Template (Steinrieselstrasse-Daten) ===
        "Wand- und Bodenbeläge UG/EG/1.OG": data.get("baustelle_objekt", ""),
        "Gartenstrasse 14": data.get("baustelle_adresse", ""),
        "317 Neuenegg": data.get("baustelle_plz_ort", ""),
        "Bautox GmbH": data.get("auftraggeber_name", ""),
        "Alpenstrasse 50": data.get("auftraggeber_adresse", ""),
        "3052 Zollikofen": data.get("auftraggeber_plz_ort", ""),
        "Urs Schürmann": data.get("bauherr_firma", ""),
        "Steinrieselstrasse 2": data.get("bauherr_adresse", ""),
        "3203 Mühleberg": data.get("bauherr_plz_ort", ""),
        "031 751 00 73": data.get("bauherr_telefon", ""),
        "Hinterkappelen 02.08.2022": f"{data.get('ort_bearbeitung', '')} {datum_bearb}",
        "Hinterkappelen, 02.08.2022": f"{data.get('ort_bearbeitung', '')}, {datum_bearb}",
        "Asbesthaltiger Klebemörtel auf Wandfliesen": data.get("schadstoff_art", ""),
        "68 m2": data.get("schadstoff_menge", ""),
        "15.08.2022 bis 31.08.2022": sanierung_zeitraum,
        "Besprechung/Baustellensitzung/Begehung vom 18.07.2022":
            f"Besprechung/Baustellensitzung/Begehung vom {besprech}" if besprech else
            "Besprechung/Baustellensitzung/Begehung vom",
        "Lulzim Bushi 756.5555.0911.16, \t\tJose Alberto Chica Briones 756.7932.7946.58": p_line2,

        # Ersteller (beide Templates)
        "Cintia Reitmann": data.get("ersteller_name", "Cintia Reitmann"),
    }

    replace_in_document(doc, replacements)
    doc.save(output_path)


# =========================================================================
# Luftbilanz Generator
# =========================================================================

def generate_luftbilanz(data: dict, template_path: str, output_path: str):
    """Fuellt das Luftbilanz-Excel-Template."""
    from openpyxl.styles import Alignment

    shutil.copyfile(template_path, output_path)
    os.chmod(output_path, 0o664)
    wb = openpyxl.load_workbook(output_path)
    ws = wb.active

    objekt_text = f"Asbestsanierung: {data.get('baustelle_objekt', '')}, {data.get('baustelle_adresse', '')}, {data.get('baustelle_plz_ort', '')}"
    zone_text = data.get("lb_zone", data.get("schadstoff_art", ""))

    ws["C3"] = objekt_text
    ws["C4"] = zone_text

    # Text in C3/C4 vollstaendig sichtbar machen:
    # 1) ueber mehrere Spalten verbinden (C bis F)
    # 2) Zeilenumbruch + vertikale Ausrichtung aktivieren
    # 3) Zeilenhoehe automatisch anpassen
    for row, text in [(3, objekt_text), (4, zone_text)]:
        try:
            ws.unmerge_cells(start_row=row, start_column=3, end_row=row, end_column=6)
        except Exception:
            pass
        ws.merge_cells(start_row=row, start_column=3, end_row=row, end_column=6)
        cell = ws.cell(row=row, column=3)
        cell.alignment = Alignment(wrap_text=True, vertical="center", horizontal="left")
        # Zeilenhoehe: ca. 15pt pro Zeile, schaetzen anhand Textlaenge
        approx_chars_per_line = 60
        lines = max(1, (len(text) // approx_chars_per_line) + (1 if len(text) % approx_chars_per_line else 0))
        ws.row_dimensions[row].height = max(20, lines * 16)

    # Spalte C breiter, damit auch ohne Merge der Text gut sichtbar ist
    ws.column_dimensions["C"].width = max(ws.column_dimensions["C"].width or 0, 35)

    try:
        ws["B9"] = float(data.get("lb_laenge", 0))
        ws["C9"] = float(data.get("lb_breite", 0))
        ws["D9"] = float(data.get("lb_hoehe", 0))
    except (TypeError, ValueError):
        pass

    try:
        ws["E20"] = int(data.get("lb_luftwechsel", 10))
    except (TypeError, ValueError):
        ws["E20"] = 10

    wb.save(output_path)


# =========================================================================
# Alarmliste Generator
# =========================================================================

def _alarm_header_cell(cell, text: str):
    """Rote Header-Zelle im Alarmlisten-Stil."""
    set_cell_bg(cell, "C8102E")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(11)
    run.font.name = "Arial"


def _alarm_value_cell(cell, lines):
    """Weisse Wertzelle mit einer oder mehreren Textzeilen."""
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    if isinstance(lines, str):
        lines = [lines]
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(11)
        run.font.name = "Arial"


def generate_alarmliste(data: dict, output_path: str):
    """Erstellt die Alarmliste fuer Baustellen im R+S-Stil."""
    doc = Document()

    # A4-Format (210 x 297 mm) und Seitenraender
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Titel - grosser roter Balken
    title_table = doc.add_table(rows=1, cols=1)
    cell = title_table.cell(0, 0)
    set_cell_bg(cell, "C8102E")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ALARMLISTE FÜR BAUSTELLEN")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "Arial"

    doc.add_paragraph()

    # Jeder Block in eigener Tabelle -> weisser Abstand dazwischen
    blocks = [
        ("BAUSTELLEN-BEZEICHNUNG:",
            [data.get("baustelle_objekt", "")]),
        ("ADRESSE:",
            [data.get("baustelle_adresse", ""), data.get("baustelle_plz_ort", "")]),
        ("KOORDINATEN REGA:",
            [data.get("rega_koordinaten", "")]),
        ("NÄCHSTER ARZT:",
            [data.get("arzt_name", ""),
             data.get("arzt_adresse", ""),
             data.get("arzt_plz_ort", ""),
             data.get("arzt_telefon", "")]),
        ("NÄCHSTES SPITAL:",
            [data.get("spital_name", ""),
             data.get("spital_adresse", ""),
             data.get("spital_plz_ort", ""),
             data.get("spital_telefon", "")]),
    ]

    for label, values in blocks:
        block_table = doc.add_table(rows=1, cols=2)
        block_table.columns[0].width = Cm(6)
        block_table.columns[1].width = Cm(11)
        _alarm_header_cell(block_table.rows[0].cells[0], label)
        _alarm_value_cell(block_table.rows[0].cells[1], [v for v in values if v])
        block_table.rows[0].cells[0].width = Cm(6)
        block_table.rows[0].cells[1].width = Cm(11)
        # kleine weisse Luecke zwischen den Bloecken
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after = Pt(4)

    doc.add_paragraph()

    # NOTRUF-Tabelle
    notruf_table = doc.add_table(rows=5, cols=3)
    notruf_items = [
        ("RETTUNGSDIENST:", "144"),
        ("FEUERWEHR:", "118"),
        ("POLIZEI:", "117"),
        ("VERGIFTUNG:", "145"),
        ("REGA:", "1414"),
    ]

    # Linke Zelle (Merged): "NOTRUF:" (rot, vertikal zentriert)
    notruf_table.rows[0].cells[0].merge(notruf_table.rows[4].cells[0])
    _alarm_header_cell(notruf_table.rows[0].cells[0], "NOTRUF:")
    notruf_table.rows[0].cells[0].width = Cm(6)

    for i, (label, number) in enumerate(notruf_items):
        _alarm_value_cell(notruf_table.rows[i].cells[1], label)
        _alarm_value_cell(notruf_table.rows[i].cells[2], number)
        notruf_table.rows[i].cells[1].width = Cm(6)
        notruf_table.rows[i].cells[2].width = Cm(5)

    doc.add_paragraph()

    # Benachrichtigung bei Unfall
    benach_table = doc.add_table(rows=1, cols=3)
    _alarm_header_cell(benach_table.cell(0, 0), "BENACHRICHTIGUNG BEI UNFALL:")
    _alarm_value_cell(benach_table.cell(0, 1), "Gruppenführer:")
    _alarm_value_cell(benach_table.cell(0, 2),
                      [data.get("gruppenfuehrer_name", ""),
                       data.get("gruppenfuehrer_telefon", "")])
    benach_table.cell(0, 0).width = Cm(6)
    benach_table.cell(0, 1).width = Cm(4)
    benach_table.cell(0, 2).width = Cm(7)

    # Seitenumbruch vor Seite 2
    doc.add_page_break()

    # Weg zum Spital - Header rot, Wert hell (weniger rote Tinte)
    weg_table = doc.add_table(rows=1, cols=1)
    cell = weg_table.cell(0, 0)
    set_cell_bg(cell, "C8102E")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("WEG ZUM NÄCHSTEN SPITAL:")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = "Arial"

    # weisser Abstand zwischen Header und Wert
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)

    weg_val_table = doc.add_table(rows=1, cols=1)
    cell = weg_val_table.cell(0, 0)
    set_cell_bg(cell, "F5F5F5")
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(data.get("weg_zum_spital", ""))
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    run.font.name = "Arial"

    # Platzhalter fuer Karte
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[ Hier Screenshot der Anfahrtsroute einfügen ]")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()

    # Erstellt / Eingesehen
    datum_bearb = format_date(data.get("datum_bearbeitung"))

    sig_table = doc.add_table(rows=3, cols=2)
    # Graue Header fuer Unterschriftsfelder (weniger rote Tinte)
    for cell, label in [(sig_table.cell(0, 0), "ERSTELLT:"),
                        (sig_table.cell(0, 1), "EINGESEHEN:")]:
        set_cell_bg(cell, "D9D9D9")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(label)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        run.font.size = Pt(11)
        run.font.name = "Arial"

    _alarm_value_cell(sig_table.cell(1, 0), "(Datum, Bauführer, Unterschrift)")
    _alarm_value_cell(sig_table.cell(1, 1), "(Datum, Polier, Unterschrift)")

    _alarm_value_cell(sig_table.cell(2, 0),
                      [datum_bearb, data.get("ersteller_name", "")])
    _alarm_value_cell(sig_table.cell(2, 1),
                      [data.get("gruppenfuehrer_name", "")])

    doc.save(output_path)


# =========================================================================
# Alles generieren
# =========================================================================

def generate_all(data: dict, output_dir: str, template_dir: str) -> dict:
    """Generiert alle 3 Dokumente (Sanierungsplan, Luftbilanz, Alarmliste)."""
    os.makedirs(output_dir, exist_ok=True)

    # Welche Sanierungsplan-Variante? default: ekas
    variant = data.get("sanierungsplan_variante", "ekas")
    if variant == "factsheet":
        sanierungsplan_template = os.path.join(template_dir, "sanierungsplan_factsheet_template.docx")
    else:
        sanierungsplan_template = os.path.join(template_dir, "sanierungsplan_ekas_template.docx")

    paths = {
        "sanierungsplan": os.path.join(output_dir, "Sanierungsplan.docx"),
        "luftbilanz": os.path.join(output_dir, "Luftbilanz.xlsx"),
        "alarmliste": os.path.join(output_dir, "Alarmliste.docx"),
    }

    generate_sanierungsplan(data, sanierungsplan_template, paths["sanierungsplan"])
    generate_luftbilanz(
        data,
        os.path.join(template_dir, "luftbilanz_template.xlsx"),
        paths["luftbilanz"],
    )
    generate_alarmliste(data, paths["alarmliste"])

    return paths
